import os
import re
import json
import time
import fitz
import hashlib
import pandas as pd
import google.generativeai as genai
import streamlit as st

from io import BytesIO
from docx import Document
from dotenv import load_dotenv

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or st.secrets.get("GEMINI_API_KEY", None)


# =========================
# RESUME PARSER
# =========================
def extract_text_from_pdf(pdf_file):
    text = ""
    try:
        if hasattr(pdf_file, "read"):
            pdf_file.seek(0)
            pdf_bytes = pdf_file.read()
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        else:
            doc = fitz.open(pdf_file)

        for page in doc:
            text += page.get_text()

        doc.close()
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF: {e}")


def extract_text_from_docx(docx_file):
    text = ""
    try:
        if hasattr(docx_file, "read"):
            docx_file.seek(0)
            file_bytes = docx_file.read()
            doc = Document(BytesIO(file_bytes))
        else:
            doc = Document(docx_file)

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX: {e}")


def extract_text_from_resume(resume_file):
    file_name = ""

    if hasattr(resume_file, "name"):
        file_name = resume_file.name.lower()
    elif isinstance(resume_file, str):
        file_name = resume_file.lower()

    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(resume_file)
    elif file_name.endswith(".docx"):
        return extract_text_from_docx(resume_file)
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")


# =========================
# HELPERS
# =========================
def get_file_hash(uploaded_file):
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    return hashlib.md5(file_bytes).hexdigest()


def extract_json_from_text(text):
    text = text.strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    match = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            pass

    match = re.search(r"(\{.*\}|\[.*\])", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            pass

    raise ValueError("Could not parse JSON from model response")


# =========================
# GEMINI
# =========================
def setup_gemini(model_name="gemini-1.5-flash"):
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not found in .env file or Streamlit secrets")

    genai.configure(api_key=GEMINI_API_KEY)
    return genai.GenerativeModel(model_name)


def safe_generate_content(prompt, model_names=None, max_retries=3):
    """
    Safe Gemini call with:
    - retry on 429 / temporary failure
    - fallback model support
    """
    if model_names is None:
        model_names = ["gemini-1.5-flash", "gemini-2.5-flash"]

    last_error = None

    for model_name in model_names:
        try:
            model = setup_gemini(model_name)
        except Exception as e:
            last_error = e
            continue

        for attempt in range(max_retries):
            try:
                response = model.generate_content(prompt)
                if response and hasattr(response, "text") and response.text:
                    return response.text.strip()
                raise Exception("Empty response from Gemini")
            except Exception as e:
                last_error = e
                err = str(e).lower()

                # quota / rate limit / temporary retry
                if "429" in err or "quota" in err or "rate limit" in err:
                    wait_time = min(50, 10 * (attempt + 1))
                    time.sleep(wait_time)
                    continue

                if "503" in err or "timeout" in err or "temporarily unavailable" in err:
                    wait_time = min(20, 5 * (attempt + 1))
                    time.sleep(wait_time)
                    continue

                # non-retryable
                break

    raise Exception(f"Gemini API failed: {last_error}")


# =========================
# SKILL EXTRACTION
# =========================
@st.cache_data(show_spinner=False)
def extract_skills_with_gemini_cached(resume_text):
    prompt = f"""
Extract only professional and technical skills from this resume.

Rules:
- Return JSON only
- Format: {{"skills": ["skill1", "skill2", "skill3"]}}
- Include programming languages, frameworks, libraries, tools, platforms, databases, cloud, analytics tools
- Remove duplicates
- Keep names short and standard
- Do not merge all skills into one string
- Do not include soft skills
- Do not include education, project names, or certifications unless they are technologies

Resume:
{resume_text[:15000]}
"""

    response_text = safe_generate_content(
        prompt=prompt,
        model_names=["gemini-1.5-flash", "gemini-2.5-flash"],
        max_retries=3
    )

    data = extract_json_from_text(response_text)

    skills = data.get("skills", [])
    cleaned_skills = []

    for skill in skills:
        skill = str(skill).strip()
        if not skill:
            continue

        parts = re.split(r",|\n|\||•|;", skill)
        for part in parts:
            part = part.strip()
            if part:
                cleaned_skills.append(part)

    final_skills = []
    known_tokens = [
        "Python", "Java", "JavaScript", "TypeScript", "HTML", "CSS",
        "React", "Next.js", "Vue.js", "Node.js", "Express", "FastAPI", "Flask",
        "MongoDB", "MySQL", "PostgreSQL", "Redis", "SQL", "GraphQL", "REST APIs",
        "AWS", "Google Cloud", "Azure", "Docker", "Kubernetes", "Git", "GitHub Actions",
        "CI/CD", "Apache Kafka", "Microservices", "System Architecture", "Tailwind CSS",
        "Power BI", "Excel", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "PyTorch",
        "Machine Learning", "Deep Learning", "NLP", "LLM", "RAG", "Linux", "Go", "PHP",
        "Agents", "MLOps", "Deployment"
    ]

    for item in cleaned_skills:
        matched_any = False
        for token in known_tokens:
            if token.lower() in item.lower():
                final_skills.append(token)
                matched_any = True
        if not matched_any:
            final_skills.append(item)

    return sorted(list(set([x.strip() for x in final_skills if x.strip()])))


# =========================
# LOCAL ROLE SKILL MAP
# =========================
LOCAL_ROLE_SKILL_MAP = {
    "Backend Developer": [
        "Python", "Java", "Go", "Node.js", "Express", "FastAPI", "Flask",
        "REST APIs", "GraphQL", "MongoDB", "MySQL", "PostgreSQL", "Redis",
        "Docker", "Kubernetes", "Git", "CI/CD", "Microservices", "Apache Kafka"
    ],
    "Frontend Developer": [
        "HTML", "CSS", "JavaScript", "TypeScript", "React", "Next.js",
        "Vue.js", "REST APIs", "Git", "Tailwind CSS"
    ],
    "Full Stack Developer": [
        "HTML", "CSS", "JavaScript", "TypeScript", "React", "Next.js",
        "Vue.js", "Node.js", "Express", "FastAPI", "Flask", "REST APIs",
        "GraphQL", "MongoDB", "MySQL", "PostgreSQL", "Git", "Docker", "Redis"
    ],
    "Software Engineer": [
        "Python", "Java", "Go", "JavaScript", "TypeScript", "SQL",
        "Git", "Docker", "REST APIs", "System Architecture", "CI/CD",
        "Microservices"
    ],
    "Python Developer": [
        "Python", "FastAPI", "Flask", "SQL", "MySQL",
        "PostgreSQL", "REST APIs", "Git", "Docker"
    ],
    "Node.js Developer": [
        "JavaScript", "TypeScript", "Node.js", "Express", "MongoDB",
        "MySQL", "PostgreSQL", "REST APIs", "GraphQL", "Git", "Docker", "Redis"
    ],
    "React Developer": [
        "HTML", "CSS", "JavaScript", "TypeScript", "React",
        "Next.js", "REST APIs", "Git", "Tailwind CSS"
    ],
    "DevOps Engineer": [
        "AWS", "Google Cloud", "Azure", "Docker", "Kubernetes",
        "CI/CD", "GitHub Actions", "Linux", "Microservices"
    ],
    "Cloud Engineer": [
        "AWS", "Google Cloud", "Azure", "Docker", "Kubernetes",
        "Linux", "CI/CD", "System Architecture"
    ],
    "Web Developer": [
        "HTML", "CSS", "JavaScript", "React", "Node.js",
        "Express", "MongoDB", "REST APIs", "Git"
    ],
    "Data Analyst": [
        "Python", "SQL", "Excel", "Power BI", "Pandas", "NumPy"
    ],
    "Data Scientist": [
        "Python", "SQL", "Pandas", "NumPy", "Scikit-learn",
        "Machine Learning", "Deep Learning"
    ],
    "AI/ML Engineer": [
        "Python", "SQL", "Pandas", "NumPy", "Scikit-learn",
        "Machine Learning", "Deep Learning", "RAG", "Agents", "MLOps", "Deployment"
    ]
}


# =========================
# NORMALIZATION
# =========================
def normalize_skill(skill):
    skill = str(skill).lower().strip()

    replacements = {
        "artificial intelligence": "ai",
        "machine learning": "ml",
        "deep learning": "dl",
        "natural language processing": "nlp",
        "large language model": "llm",
        "large language models": "llm",
        "retrieval augmented generation": "rag",
        "structured query language": "sql",
        "sql database": "sql",
        "python programming": "python",
        "scikit learn": "scikit-learn",
        "nodejs": "node.js",
        "node js": "node.js",
        "nextjs": "next.js",
        "github actions": "ci/cd",
        "rest api": "rest apis",
        "restful api": "rest apis",
        "restful apis": "rest apis",
        "google cloud platform": "google cloud",
        "gcp": "google cloud",
        "amazon web services": "aws",
        "postgres": "postgresql",
        "mlops/deployment": "mlops deployment",
        "deployment": "mlops deployment",
        "mlops": "mlops deployment",
        "agentic ai": "agents",
        "ai agents": "agents"
    }

    skill = re.sub(r"[\(\)]", " ", skill)
    skill = re.sub(r"[^a-zA-Z0-9+#\-\. /]", " ", skill)
    skill = re.sub(r"\s+", " ", skill).strip()

    return replacements.get(skill, skill)


def calculate_match_and_missing(extracted_skills, required_skills):
    user_norm_map = {}
    for skill in extracted_skills:
        norm = normalize_skill(skill)
        if norm:
            user_norm_map[norm] = skill

    required_clean = []
    seen = set()

    for skill in required_skills:
        clean = normalize_skill(skill)
        if clean and clean not in seen:
            seen.add(clean)
            required_clean.append(skill)

    matched = []
    missing = []

    for skill in required_clean:
        required_norm = normalize_skill(skill)

        if required_norm in user_norm_map:
            matched.append(skill)
            continue

        found = False
        for user_norm in user_norm_map.keys():
            if required_norm in user_norm or user_norm in required_norm:
                matched.append(skill)
                found = True
                break

        if not found:
            missing.append(skill)

    total = len(required_clean)
    score = round((len(matched) / total) * 100, 2) if total > 0 else 0.0

    return score, matched, missing


# =========================
# COURSES FROM CSV ONLY
# =========================
@st.cache_data
def load_courses_csv(csv_path="Courses.csv"):
    if not os.path.exists(csv_path):
        if os.path.exists("courses.csv"):
            csv_path = "courses.csv"
        else:
            return pd.DataFrame(columns=["skill", "course_name"])

    df = pd.read_csv(csv_path)

    df.columns = [str(col).strip().lower() for col in df.columns]

    required_cols = {"skill", "course_name"}
    if not required_cols.issubset(set(df.columns)):
        raise ValueError("Courses CSV must contain these columns: skill, course_name")

    df = df[["skill", "course_name"]].copy()
    df["skill"] = df["skill"].astype(str).str.strip()
    df["course_name"] = df["course_name"].astype(str).str.strip()

    df = df[(df["skill"] != "") & (df["course_name"] != "")]
    df = df.drop_duplicates().reset_index(drop=True)

    return df


def fetch_courses_for_skill(skill, courses_df, max_courses=3):
    if courses_df.empty:
        return [{
            "skill": skill,
            "course_name": f"Learn {skill}"
        }]

    skill_norm = normalize_skill(skill)
    matched_rows = []

    for _, row in courses_df.iterrows():
        csv_skill = str(row["skill"]).strip()
        csv_skill_norm = normalize_skill(csv_skill)

        if (
            skill_norm == csv_skill_norm or
            skill_norm in csv_skill_norm or
            csv_skill_norm in skill_norm
        ):
            matched_rows.append({
                "skill": skill,
                "course_name": row["course_name"]
            })

    if not matched_rows:
        return [{
            "skill": skill,
            "course_name": f"Learn {skill}"
        }]

    return matched_rows[:max_courses]


# =========================
# ROLE ANALYSIS DIRECTLY FROM SKILLS
# =========================
def build_role_wise_analysis_from_local_map(extracted_skills, courses_df):
    analysis = []

    for role, required_skills in LOCAL_ROLE_SKILL_MAP.items():
        score, matched_skills, missing_skills = calculate_match_and_missing(
            extracted_skills, required_skills
        )

        recommended_courses = []
        for skill in missing_skills[:4]:
            recommended_courses.extend(
                fetch_courses_for_skill(skill, courses_df=courses_df, max_courses=2)
            )

        role_courses_df = pd.DataFrame(recommended_courses)
        if role_courses_df.empty:
            role_courses_df = pd.DataFrame(columns=["skill", "course_name"])
        else:
            role_courses_df = role_courses_df.drop_duplicates().reset_index(drop=True)

        analysis.append({
            "job_role": role,
            "match_score": score,
            "required_skills": required_skills,
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
            "recommended_courses": role_courses_df
        })

    analysis = sorted(analysis, key=lambda x: x["match_score"], reverse=True)
    return analysis


# =========================
# CHATBOT
# =========================
def generate_chatbot_response(user_question, extracted_skills, role_analysis):
    compact_roles = []

    for item in role_analysis:
        compact_roles.append({
            "job_role": item["job_role"],
            "match_score": item["match_score"],
            "required_skills": item["required_skills"],
            "matched_skills": item["matched_skills"],
            "missing_skills": item["missing_skills"]
        })

    prompt = f"""
You are a helpful career guidance assistant.

Candidate extracted skills:
{", ".join(extracted_skills)}

Role analysis:
{json.dumps(compact_roles, ensure_ascii=False)}

User question:
{user_question}

Instructions:
- Answer clearly, practically, and simply
- Mention best matching roles first
- Mention what to learn next based on missing skills
- Keep the answer structured and concise
"""

    response_text = safe_generate_content(
        prompt=prompt,
        model_names=["gemini-1.5-flash", "gemini-2.5-flash"],
        max_retries=3
    )
    return response_text


# =========================
# MAIN PIPELINE
# =========================
def run_resume_analysis(resume_file):
    resume_text = extract_text_from_resume(resume_file)
    extracted_skills = extract_skills_with_gemini_cached(resume_text)

    courses_df = load_courses_csv("Courses.csv")

    role_analysis = build_role_wise_analysis_from_local_map(
        extracted_skills=extracted_skills,
        courses_df=courses_df
    )

    matched_jobs_df = pd.DataFrame([
        {
            "job_role": item["job_role"],
            "match_score": item["match_score"]
        }
        for item in role_analysis
    ])

    return {
        "resume_text": resume_text,
        "extracted_skills": extracted_skills,
        "fetched_jobs": [],
        "matched_jobs": matched_jobs_df,
        "role_analysis": role_analysis
    }